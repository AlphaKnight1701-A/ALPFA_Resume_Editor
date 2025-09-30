# Imports Appropriate Libraries
import pdfminer.high_level as pdf
import os

#Imports Additional Components/Classes
from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

# Creates the Flask Application
app = Flask(__name__)

# Sets the Upload Folder and Template File Designations
UPLOAD_FOLDER = "uploads"
TEMPLATE_FILE = "Resume 2025.docx"

# Connects/Verifies the Upload Folder Exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Function that Isolates
def isolate_string_by_keyword(text, keyword):

    # 
    """
    Scans text for a keyword and returns the complete string containing that keyword,
    bounded by spaces before and after.
    """

    # Splits the Text into an Array of Lines
    lines = text.split('\n')
    for line in lines:
        if keyword in line:
            words = line.split()
            for word in words:
                if keyword in word:
                    return word.strip()
    return ""

# Function that Extracts Text from a PDF or DOCX File
def extract_text(file_path):
    
    # Condition Statement that Checks if the File is a PDF
    if file_path.endswith(".pdf"):

        # Extracts Text from the PDF File
        return pdf.extract_text(file_path)

    # Condition Statement that Checks if the File is a DOCX
    elif file_path.endswith(".docx"):

        # Opens a new DOCX File
        doc = Document(file_path)

        # Returns Extracted Text from Word Document Using Paragraphs and Separated by NewLine Character
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    # Returns Fail Message if File Type is Invalid
    return "Unsupported file format"

# Function that Identifies and Formats Sections from Input Text Using Key Words (As Seen in Resume Template)
def extract_sections(text):

    # Defines Section Key-Word Dictionary for ALL Resume Sections
    sections = {
        "PERSONAL INFORMATION": [],
        "EDUCATION": [],
        "EXPERIENCE": [],
        "LEADERSHIP AND PROFESSIONAL DEVELOPMENT": [],
        "SKILLS, ACTIVITIES & INTERESTS": []
    }

    # Sets Current Section to None
    current_section = None

    # Separates the Block of Text into An Array of Lines
    lines = text.split("\n")

    # Loop that Iterates through each Line in the Text
    for line in lines:

        # Strips the Line of Whitespace (Cleans it Up)
        line = line.strip()

        # Condition Statement that Checks if Line Is Empty
        if not line:

            # Skips to the Next Line
            continue

        # Sets Line to Upper Case
        upper_line = line.upper()

        # Condition Statement that Checks if upper_line contains any Section Title in Section Dictionary
        if any(section in upper_line for section in sections.keys()):

            # Loop that Iterates through each Section Title in Section Dictionary
            for section in sections.keys():

                # Condition Statement that Checks if Line Contains the Section Title
                if section in upper_line:

                    # Sets Current Section to the Section Title
                    current_section = section

                    # Skips to the Next Line
                    break

        # Condition Statement that Checks if Current Section is NOT NULL
        elif current_section and line:

            # Appends the Line to the Current Section in the Sections Dictionary
            sections[current_section].append(line)

    # Debug print statements
    print("\n=== Debug: Sections Dictionary ===")
    for section, content in sections.items():
        print(f"\n{section}:")
        for item in content:
            print(f"  - {item}")
    print("\n================================")
    
    # Returns the Sections Dictionary
    return sections

# Function that Creates a Resume Template Document with extracted content
def format_resume(sections):
    
    # Loads Resume Template Document
    doc = Document("Resume 2025.docx")

    # Declares
    current_section = None
    section_content = None

    # Process each paragraph
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().upper()

        # Check if this is a section header
        for section in sections.keys():
            if section in text:
                current_section = section
                section_content = sections[section]
                break

        # If we're in a section and have content to add
        if current_section and section_content and len(section_content) > 0:
            # Skip the section header itself
            if section_content and not any(section in text for section in sections.keys()):
                # Clear existing text while preserving formatting
                for run in paragraph.runs:
                    if run.text.strip():
                        content = section_content.pop(0) if section_content else ""
                        run.text = content
                        break

    output_path = os.path.join(UPLOAD_FOLDER, "formatted_resume.docx")
    doc.save(output_path)
    return output_path

@app.route("/", methods=["GET", "POST"])
def upload_resume():
    if request.method == "POST":
        file = request.files.get("resume")
        if not file or not file.filename:
            return "No file uploaded or invalid file!"

        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)

        raw_text = extract_text(file_path)
        structured_sections = extract_sections(raw_text)
        formatted_resume_path = format_resume(structured_sections)

        return send_file(formatted_resume_path, as_attachment=True)

    return render_template("upload.html")

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, debug=True)